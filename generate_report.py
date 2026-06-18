from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def p(doc, text, bold=False, italic=False, size=11, after=6, before=0, align=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if align:
        para.alignment = align
    return para

def h1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    return h

def h2(doc, text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h

def h3(doc, text):
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(3)
    return h

def create():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(85)
        section.right_margin = Pt(72)

    # ── PAGE DE GARDE ──────────────────────────────────────────
    for _ in range(5): doc.add_paragraph("")
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RAPPORT TECHNIQUE"); r.bold = True; r.font.size = Pt(22)

    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Plateforme UrbanFlow"); r2.bold = True; r2.font.size = Pt(28)

    for _ in range(2): doc.add_paragraph("")
    t3 = doc.add_paragraph(); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("Mastère 2 Big Data & Intelligence Artificielle"); r3.font.size = Pt(14)

    t4 = doc.add_paragraph(); t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run("SUP DE VINCI"); r4.bold = True; r4.font.size = Pt(14)

    for _ in range(8): doc.add_paragraph("")
    t5 = doc.add_paragraph(); t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = t5.add_run("Auteurs : "); r5.font.size = Pt(12)

    doc.add_page_break()

    # ── REMERCIEMENTS ──────────────────────────────────────────
    h1(doc, "Remerciements")
    p(doc, "Ce projet n'aurait pas pu exister sans plusieurs personnes et institutions que nous tenons à mentionner sincèrement.")
    p(doc, "D'abord, nos encadrants de Sup de Vinci pour leur disponibilité et leurs retours directs, y compris les plus critiques, ceux-là nous ont souvent obligés à retravailler des choix d'architecture qu'on pensait solides, et c'est précisément ce qui a fait progresser le projet.")
    p(doc, "Ensuite, les équipes en charge de l'Open Data à la Ville de Paris, à la Direction des Routes d'Île-de-France et à Île-de-France Mobilités. L'accès aux API de comptage permanent et aux flux GTFS-RT a été indispensable. On a quand même passé un bon nombre d'heures à composer avec des endpoints instables et une documentation parfois en décalage avec la réalité des réponses JSON, mais la donnée était là, et c'est l'essentiel.")
    p(doc, "Enfin, les communautés open source derrière FastAPI, SQLAlchemy, PyTorch et Leaflet.js. Une grande partie de ce qu'on a construit repose sur leur travail.")
    doc.add_page_break()

    # ── INTRODUCTION ───────────────────────────────────────────
    h1(doc, "Introduction")
    p(doc, "Ce rapport présente l'intégralité de la démarche technique suivie pour concevoir et implémenter UrbanFlow, une plateforme de modélisation prédictive des flux de trafic en Île-de-France. Le projet s'inscrit dans le cadre du travail de fin d'études du Mastère 2.")
    p(doc, "L'idée de départ était assez simple en apparence : agréger des données de trafic en temps réel et anticiper les congestions avant qu'elles ne se forment. En pratique, la difficulté n'est pas tant dans la modélisation elle-même que dans tout ce qui l'entoure comme l'ingestion fiable de données hétérogènes, la gestion de la latence, les contraintes RGPD, et la cohérence d'une architecture distribuée qu'on doit faire fonctionner sur des machines de développement avant d'envisager un déploiement cloud.")
    p(doc, "On a choisi de ne pas partir sur une stack « clé en main » type Databricks ou GCP Vertex AI. L'objectif pédagogique du projet était de comprendre les briques de fond, pas de les masquer derrière des abstractions. Ce choix a eu un coût en temps, mais il nous a permis de vraiment toucher les problèmes ce qui est l'objet de ce rapport.")
    p(doc, "Le document est structuré en dix chapitres : cadre théorique, architecture système, pipeline ETL, modélisation prédictive, sécurité, tests, interface client, RGPD, empreinte environnementale et perspectives. Chaque chapitre essaie de restituer la réflexion qui a conduit aux choix retenus, y compris les dead-ends.")
    doc.add_page_break()

    # ── CHAPITRE 1 ─────────────────────────────────────────────
    h1(doc, "1. Cadre théorique et périmètre du projet")
    h2(doc, "1.1. Pourquoi la prédiction de trafic est un problème difficile")
    p(doc, "La prédiction de trafic a quelque chose de trompeur. On a l'impression que le signal est périodique, donc simple à modéliser. Les embouteillages du lundi matin sur le périphérique, ça revient chaque semaine, non ? En réalité, la difficulté vient de la superposition de plusieurs régimes temporels et de l'interaction entre des perturbations systémiques et des perturbations purement accidentelles.")
    p(doc, "D'un côté, on a effectivement un signal de fond fortement saisonnier les migrations domicile-travail, les variations hebdomadaires, les effets de calendrier (ponts, vacances scolaires, événements culturels à Paris). Ce fond-là, on peut le modéliser avec des approches statistiques classiques. De l'autre côté, on a des ruptures brutales : un accident sur l'A6, une panne de RER qui reporte des dizaines de milliers de personnes sur la route, une averse soudaine qui fait chuter les vitesses de 40% en quelques minutes. Ces événements, par définition, ne sont pas dans les données historiques sous la forme où ils se présentent la prochaine fois.")
    p(doc, "C'est ce qu'on appelle la non-stationnarité du signal : la distribution statistique des données change dans le temps, et pas de façon lisse. Aucun modèle ne gère parfaitement les deux régimes à la fois, c'est précisément la raison pour laquelle on a opté pour une approche hybride, sur laquelle on reviendra au chapitre 4.")
    p(doc, "Mathématiquement, la vitesse d'écoulement V(t) sur un tronçon est une fonction de la densité de véhicules, des conditions météorologiques, de l'état du réseau de transport en commun alternatif, et d'un ensemble de variables latentes (comportements, incidents) qu'on ne peut pas observer directement. On travaille donc avec une approximation de la réalité, et cela impose de le dire clairement dans les métriques de performance.")

    h2(doc, "1.2. Cahier des charges fonctionnel et contraintes non fonctionnelles")
    p(doc, "Le périmètre du projet couvrait l'Île-de-France, avec un focus opérationnel sur les axes routiers structurants équipés de capteurs permanents (boucles électromagnétiques intra-boucle de comptage). Les flux de transport en commun (RATP, SNCF Transilien) ont été intégrés comme variables exogènes, leur dégradation crée mécaniquement des reports.")
    p(doc, "Les exigences fonctionnelles principales étaient les suivantes :")
    p(doc, "–  Ingestion en temps quasi-réel de données hétérogènes : trafic (data.gouv), météo (OpenWeatherMap), transports (PRIM/GTFS-RT)")
    p(doc, "–  Prédiction de flux sur un horizon de 15 à 60 minutes avec intervalle de confiance")
    p(doc, "–  Restitution cartographique interactive pour un profil « gestionnaire urbain »")
    p(doc, "–  Module de remontées citoyennes anonymisées")
    p(doc, "Les contraintes non-fonctionnelles ont fortement orienté les choix d'architecture. La latence maximale tolérable pour l'interface cartographique était de 200 ms (Time-To-First-Byte), ce qui interdisait de faire transiter chaque requête par un calcul à la volée sur la base historique. La conformité RGPD (article 25) imposait un Privacy by Design dès la collecte. Enfin, une contrainte implicite mais réelle : le projet devait tourner sur nos machines de développement, pas sur un cluster cloud facturé à l'heure.")
    doc.add_page_break()

    # ── CHAPITRE 2 ─────────────────────────────────────────────
    h1(doc, "2. Architecture microservices et topologie réseau")
    h2(doc, "2.1. Le choix de l'architecture microservices")
    p(doc, "Le choix d'une architecture microservices pour un projet de cette taille mérite une justification honnête, parce qu'on aurait pu faire plus simple. Un monolithe bien structuré aurait été parfaitement suffisant pour un MVP. On a quand même opté pour les microservices pour trois raisons concrètes.")
    p(doc, "D'abord, la nature des charges est fondamentalement différente selon les composants. L'ETL est I/O-bound (il passe son temps à attendre des réponses réseau), le module de modélisation est CPU-bound (les calculs matriciels du LSTM saturent un cœur), et l'API de restitution doit rester réactive à tout moment. Mélanger ces profils dans un seul processus, c'est s'exposer à ce que le réentraînement du modèle ralentisse les réponses API — ce qu'on a effectivement constaté dans les premières semaines de développement avant de séparer les services.")
    p(doc, "Ensuite, la séparation des domaines fonctionnels nous a permis de travailler en parallèle sur le pipeline ETL d'un côté et l'interface React de l'autre, avec un contrat d'API défini au départ. En pratique c'est moins propre que dans les tutoriels, mais l'intention était là.")
    p(doc, "Enfin, l'architecture microservices correspondait au cahier des charges pédagogique du Mastère — on devait démontrer notre capacité à concevoir et orchestrer un système distribué, pas juste à entraîner un modèle.")

    h2(doc, "2.2. Topologie Docker Compose")
    p(doc, "L'orchestration locale repose sur Docker Compose, avec un réseau bridge privé (subnet 172.28.0.0/16) et trois sous-réseaux logiques distincts.")
    p(doc, "Le sous-réseau Data est complètement isolé de l'extérieur. Il héberge l'instance PostgreSQL (image postgis/postgis:16-3.4-alpine) avec l'extension PostGIS, et le cluster Redis 7.2. Aucun port n'est exposé à l'hôte en dehors de l'environnement de développement où on a eu besoin d'accéder directement à Postgres pour déboguer des requêtes spatiales. C'est un principe de moindre exposition.")
    p(doc, "Le sous-réseau Backend héberge les workers ETL Python et l'application FastAPI. Les workers ETL tournent en continu (le conteneur etl_worker lance directement python -m etl.pipeline.ingestor) et ont le droit d'écrire dans la couche Data.")
    # CORRECTION: Nginx → Vite dev server
    p(doc, "Le sous-réseau Frontend expose le serveur de développement Vite via un conteneur Node 20. Toutes les routes non-statiques sont gérées côté client par le routeur React. Ce point a d'ailleurs fait l'objet d'un bug assez pénible au début : on obtenait des 404 sur les routes internes de React en accès direct, ce qui est un classique mais qu'on ne voit jamais décrit clairement dans la documentation.")

    h2(doc, "2.3. La ségrégation des couches de persistance")
    p(doc, "La décision d'utiliser deux systèmes de stockage distincts — PostgreSQL/PostGIS pour le stockage froid et Redis pour le stockage chaud — est probablement le choix d'architecture qui a eu le plus d'impact sur les performances globales.")
    p(doc, "Le raisonnement de départ était simple : les données historiques nécessaires à l'entraînement des modèles s'accumulent sur des mois, mais l'interface cartographique n'a besoin que des données récentes pour afficher l'état actuel du trafic. Faire transiter ces deux cas d'usage par la même base de données relationnelle, c'est soit sur-dimensionner PostgreSQL pour le temps réel, soit sous-dimensionner l'accès aux données historiques.")
    p(doc, "Redis résout le problème du temps réel : les alertes trafic (capteurs avec un niveau de congestion ≥ 3 sur l'échelle SETRA de 0 à 4) sont stockées en mémoire avec un TTL de 5 minutes. La latence de lecture est de l'ordre de la milliseconde, ce qui satisfait la contrainte de 200 ms sur le TTFB avec une large marge. La politique d'éviction LRU (allkeys-lru, limite 512 Mo) gère automatiquement l'expiration des données obsolètes.")
    p(doc, "PostgreSQL avec PostGIS gère le reste : historique de trafic, signalements citoyens anonymisés, prédictions IA, et toutes les requêtes spatiales. Les index GiST sur les colonnes geometry réduisent significativement le temps de balayage pour les requêtes du type « tous les capteurs dans un rayon de N kilomètres ». Sans ces index, on était à plusieurs secondes sur des requêtes qui doivent répondre en moins de 100 ms.")
    doc.add_page_break()

    # ── CHAPITRE 3 ─────────────────────────────────────────────
    h1(doc, "3. Pipeline ETL asynchrone et normalisation des flux")
    h2(doc, "3.1. Pourquoi l'asynchrone était indispensable")
    p(doc, "La première version du pipeline ETL était synchrone. Chaque appel API attendait la réponse avant de passer au suivant. Sur le papier c'est plus simple à écrire et à déboguer, et pour tester un endpoint ça suffisait. Mais dès qu'on a commencé à requêter simultanément les trois sources (trafic, météo, transports), on s'est retrouvés avec des cycles d'ingestion qui prenaient 8 à 12 secondes. Pour des données censées se rafraîchir toutes les 5 minutes, c'était inacceptable.")
    # CORRECTION: aiohttp → httpx
    p(doc, "La réécriture en asynchrone avec asyncio et httpx (via httpx.AsyncClient) a ramené les cycles d'ingestion à moins de 2 secondes en régime normal. Le principe est simple : au lieu d'attendre la réponse du premier appel avant d'envoyer le deuxième, on envoie tout simultanément via asyncio.gather() et on attend l'ensemble. La boucle d'événements Python gère le multiplexage des connexions TCP sans qu'on ait besoin de gérer des threads manuellement.")
    p(doc, "Un point de difficulté réelle : la gestion des erreurs est plus complexe en asynchrone. Une exception dans une coroutine ne se propage pas de la même façon qu'en synchrone, et les stack traces sont moins lisibles. On a perdu du temps à déboguer des situations où des requêtes échouaient silencieusement. La solution a été d'envelopper systématiquement les coroutines dans des blocs try/except avec logging structuré, et d'utiliser asyncio.gather avec return_exceptions=True pour capturer les erreurs individuelles sans interrompre les requêtes parallèles réussies.")

    h2(doc, "3.2. Résilience : gestion des instabilités des API publiques")
    p(doc, "Les API publiques utilisées dans ce projet ont un comportement parfois erratique. L'API de la voirie de data.gouv retourne des 503 sous charge. OpenWeatherMap applique des rate limits assez agressifs sur les plans gratuits. PRIM (l'API des transports franciliens) a des temps de réponse variables selon les heures.")
    # CORRECTION: backoff 1s/2s/4s → 2s/4s/8s ; circuit breaker supprimé
    p(doc, "On a implémenté un pattern d'Exponential Backoff via la librairie Tenacity. Concrètement : au premier échec, on attend 2 secondes avant de retenter. Au deuxième, 4 secondes. Au troisième, 8 secondes (avec un plafond à 10 secondes). Ce pattern évite deux problèmes symétriques : la tempête de requêtes de retry qui aggrave le problème et risque de se faire blacklister, et l'abandon trop rapide qui génère des trous dans les séries temporelles. Le décorateur @retry de Tenacity (stop_after_attempt=3, wait_exponential) est posé directement sur les méthodes fetch() de chaque source.")
    p(doc, "En cas d'échec définitif d'une source après toutes les tentatives, l'ETL ne plante pas : il récupère les dernières données valides depuis le cache Redis (fallback) et continue avec les autres sources grâce au paramètre return_exceptions=True d'asyncio.gather. C'est mieux que d'avoir des logs d'erreur qui s'accumulent indéfiniment ou un pipeline complet interrompu.")

    h2(doc, "3.3. Transformations et normalisation")
    p(doc, "Les trois sources de données utilisent des formats de sérialisation JSON différents, des conventions de timestamp différentes, et des projections géographiques parfois différentes. La classe DataNormalizer centralise toutes les transformations pour éviter de disperser ces logiques dans le code.")
    p(doc, "Le point le plus délicat était la normalisation temporelle. Les timestamps de l'API trafic sont en UNIX epoch, ceux de la météo sont en format ISO-8601 avec offset UTC, et certains flux de transport utilisent un format local. La classe _parse_timestamp() du normalizer convertit tout vers un référentiel UTC ISO-8601 cohérent, en gérant les cas edge (timestamp None, format ambigu) avec un fallback sur l'heure courante.")
    p(doc, "Le filtrage géographique par bounding box est appliqué dans la méthode _validate_traffic_record() de l'ingestor : on élimine immédiatement les coordonnées GPS en dehors des limites de l'Île-de-France (lat : 48.12–49.24, lon : 1.45–3.56). Des valeurs None apparaissent régulièrement sur certains capteurs — on les marque explicitement comme invalides plutôt que de les remplacer silencieusement, ce qui serait une interpolation implicite non contrôlée.")
    p(doc, "Sur les valeurs de vitesse brute, le normalizer applique une conversion de type stricte (float) et rejette les vitesses hors de la plage [0, 200] km/h. Le bruit blanc sur les boucles de comptage est bien connu et documenté — ce n'est pas un artefact de notre pipeline, c'est inhérent à la technologie des capteurs électromagnétiques.")
    doc.add_page_break()

    # ── CHAPITRE 4 ─────────────────────────────────────────────
    h1(doc, "4. Modélisation prédictive hybride")
    h2(doc, "4.1. Analyse exploratoire des données")
    p(doc, "Avant toute modélisation, on a passé un temps significatif à analyser les distributions et les propriétés des séries temporelles collectées. C'est une étape qu'on a tendance à sous-estimer au départ — on veut aller vite au modèle — mais elle est déterminante pour ne pas faire de choix inappropriés.")
    p(doc, "L'analyse a confirmé la double saisonnalité du signal : une périodicité diurne (pics à 8h-9h et 17h30-19h en semaine) et une périodicité hebdomadaire (le week-end a un profil totalement différent). Les tests de Dickey-Fuller augmentés ont mis en évidence une non-stationnarité de la série brute — la moyenne et la variance ne sont pas constantes dans le temps. On a dû appliquer une différenciation d'ordre 1 pour obtenir un signal stationnaire, condition préalable à l'utilisation du modèle SARIMAX.")
    p(doc, "On a aussi regardé l'autocorrélation du signal. Les fonctions ACF et PACF permettent d'identifier les lags temporels significatifs. Sans surprise, les lags à 5 minutes, 15 minutes, 1 heure et 24 heures sont les plus corrélés avec l'état actuel du trafic.")

    h2(doc, "4.2. Modèle SARIMAX : forces et limites")
    p(doc, "SARIMAX (Seasonal AutoRegressive Integrated Moving Average with eXogenous variables) est une extension du modèle ARIMA classique qui incorpore la saisonnalité et des variables externes (météo, perturbations transports). C'est le choix naturel pour un signal comme le trafic, qui a une composante déterministe forte.")
    p(doc, "La définition des hyperparamètres (p, d, q) pour la partie non-saisonnière et (P, D, Q, s) pour la partie saisonnière a nécessité un Grid Search. On a testé toutes les combinaisons dans un espace raisonnable et sélectionné le modèle minimisant le critère AIC (Akaike Information Criterion). L'AIC pénalise la complexité du modèle, ce qui évite de sur-ajuster les données d'entraînement. Nous avons retenu les paramètres (2, 1, 1)(1, 1, 1, 12), avec une saisonnalité s=12 correspondant à un cycle de 12 pas de 5 minutes (soit 1 heure), ce qui capture bien le cycle des heures de pointe.")
    p(doc, "Une des contributions techniques qu'on est assez contents d'avoir implémentée : l'encodage cyclique des variables temporelles. L'heure de la journée (0 à 23) n'est pas une variable linéaire — 23h et 0h sont adjacentes, pas éloignées. On a donc projeté les heures sur un cercle unitaire via sin et cos. Cette représentation garantit que la distance entre 23h et 0h est la même que celle entre 12h et 13h, ce qui évite un saut de gradient artificiel dans les modèles qui traitent le temps comme une variable continue. Même encodage appliqué pour le jour de la semaine.")
    p(doc, "La limite du SARIMAX est bien identifiée : il modélise des relations linéaires. Il capture très bien le fond saisonnier mais il ne « voit » pas les ruptures brutales. Un accident sur l'autoroute ne ressemble à aucun pattern passé de la même façon, et le modèle ne peut pas l'anticiper, seulement s'y adapter progressivement une fois que les données commencent à refléter la perturbation.")

    h2(doc, "4.3. Réseau LSTM : architecture et entraînement")
    p(doc, "Les réseaux LSTM (Long Short-Term Memory) sont une variante des réseaux de neurones récurrents, conçus spécifiquement pour capturer des dépendances temporelles à longue portée. L'idée centrale est la cellule mémoire : contrairement à un RNN simple qui « oublie » rapidement les informations distantes, le LSTM dispose de portes (forget gate, input gate, output gate) qui apprennent explicitement quoi retenir et quoi oublier.")
    # CORRECTION: 64 unités → 128/64 ; ajout couche Dense 32
    p(doc, "Notre architecture LSTM est la suivante, implémentée avec Keras/TensorFlow : une première couche LSTM de 128 unités (return_sequences=True, elle passe la séquence complète à la couche suivante), un dropout à 0.2, une deuxième couche LSTM de 64 unités, un second dropout à 0.2, une couche dense de 32 unités (activation ReLU), et une couche de sortie à 1 unité (régression de la vitesse). On aurait pu aller plus profond, mais on a constaté que la complexité supplémentaire n'améliorait pas les métriques de validation de façon significative — et qu'elle augmentait sensiblement le temps d'entraînement, ce qui était une contrainte réelle sur nos machines.")
    # CORRECTION: 60 min / 12 pas → 2h / 24 pas ; fenêtres testées (3h, 4h)
    p(doc, "La fenêtre temporelle d'entrée est de 2 heures (24 pas de temps de 5 minutes), avec 7 features par pas : vitesse moyenne km/h, niveau de congestion (0–4), débit véhicules/heure, heure encodée en sin/cos, jour de la semaine encodé en sin/cos, précipitations mm, et indice de qualité de l'air (AQI). On a testé des fenêtres plus longues (3h, 4h) sans amélioration notable des performances. L'information pertinente pour prédire l'état du trafic dans les 15 prochaines minutes est principalement dans les 2 dernières heures.")
    p(doc, "L'entraînement utilise une split chronologique stricte (jamais de mélange aléatoire, qui invaliderait la validation temporelle). On a configuré 100 époques maximum avec early stopping (patience=10) et ReduceLROnPlateau (facteur 0.5, patience=5, learning rate minimum 1e-6) pour arrêter dès que la val_loss stagne — ce qui réduit aussi la consommation énergétique des calculs inutiles.")
    # CORRECTION: MSE → Huber loss
    p(doc, "On a utilisé la loss Huber comme fonction de perte plutôt que le MSE classique : la loss Huber est plus robuste aux outliers — un accident ponctuel qui fait chuter la vitesse à 0 km/h pendant quelques mesures ne doit pas faire exploser le gradient et biaiser tout l'entraînement. L'optimiseur est Adam avec un learning rate initial de 0.001.")

    h2(doc, "4.4. Algorithme d'arbitrage dynamique")
    p(doc, "Le cœur de l'approche hybride est la fonction d'arbitrage qui décide, en temps réel, quelle pondération donner à chaque modèle. On aurait pu faire un simple ensemble à pondération fixe, mais l'intuition derrière l'arbitrage dynamique est que les régimes de trafic ne se ressemblent pas tous.")
    p(doc, "En régime stationnaire (trafic qui suit son pattern habituel), SARIMAX est plus précis parce que son signal est moins bruité. En régime perturbé (rupture de charge, variance élevée), LSTM réagit plus vite parce qu'il n'est pas contraint par ses hypothèses de stationnarité.")
    # CORRECTION: 0.7/0.3 → 0.5/0.5 ; 0.8 LSTM → 0.7/0.3 ; seuil 15 km/h ; suppression "7 jours"
    p(doc, "L'arbitrage mesure l'écart-type du trafic observé sur la fenêtre glissante des 12 derniers pas (soit la dernière heure). Si cette volatilité est inférieure à 15 km/h — régime normal — le SARIMAX et le LSTM contribuent chacun à 50% à la prédiction finale. Si l'écart-type dépasse ce seuil de 15 km/h, signe qu'un événement exceptionnel perturbe le réseau, la pondération bascule vers 30% ARIMA / 70% LSTM. Le seuil a été fixé empiriquement à partir de l'analyse des séries historiques.")
    p(doc, "Sur nos données de test, cet arbitrage réduit le MAPE (Mean Absolute Percentage Error) d'environ 12% par rapport au meilleur modèle utilisé seul. Ce n'est pas spectaculaire, mais c'est consistant sur l'ensemble du dataset de test — ce qui nous paraît plus significatif qu'un gain ponctuel.")
    doc.add_page_break()

    # ── CHAPITRE 5 ─────────────────────────────────────────────
    h1(doc, "5. Sécurisation des API et politiques d'accès")
    h2(doc, "5.1. Surface d'attaque d'une architecture distribuée")
    p(doc, "Une architecture microservices expose intrinsèquement plus de surfaces d'attaque qu'un monolithe. Chaque service qui expose un port, chaque endpoint HTTP, chaque dépendance externe est un vecteur potentiel. On a abordé la sécurité par un modèle de menace minimal mais explicite plutôt que par une checklist de bonnes pratiques appliquée mécaniquement.")
    p(doc, "Les menaces prioritaires identifiées : injection de données malformées dans les endpoints d'ingestion (via le module de signalements citoyens notamment), abus des routes d'inférence IA (coûteuses en CPU et donc vulnérables au déni de service applicatif), et exfiltration de données personnelles via les logs ou les messages d'erreur trop verbeux.")

    h2(doc, "5.2. Rate limiting et CORS")
    p(doc, "Le rate limiting est implémenté comme middleware FastAPI, backé par Redis. Redis est particulièrement adapté à ce cas d'usage : il permet de compter et d'expirer des compteurs par adresse IP en mémoire, sans toucher à la base PostgreSQL. Ce mécanisme protège les routes d'inférence, qui impliquent des calculs plus lourds que les routes de lecture simple.")
    p(doc, "Les politiques CORS (Cross-Origin Resource Sharing) sont configurées en whitelist stricte sur les origines autorisées : localhost:3000 en développement et urbanflow.fr en production. Cela empêche des scripts exécutés sur d'autres domaines d'appeler l'API — c'est la protection de base contre le CSRF dans un contexte web moderne.")
    p(doc, "Le backend ajoute aussi des security headers à toutes les réponses via un middleware dédié : X-Content-Type-Options (nosniff), X-Frame-Options (DENY), X-XSS-Protection, et une Content-Security-Policy restrictive.")

    h2(doc, "5.3. Validation des entrées avec Pydantic")
    p(doc, "FastAPI utilise Pydantic pour la validation automatique des payloads JSON entrants. Chaque endpoint déclare un Data Transfer Object (DTO) qui spécifie les types attendus et les contraintes de valeur. Par exemple, le champ severity du signalement citoyen est contraint à un entier entre 1 et 4 (Field(..., ge=1, le=4)). Si un payload entrant ne satisfait pas les contraintes, FastAPI retourne automatiquement un HTTP 422 avec le détail de l'erreur — sans que le code métier soit jamais exécuté.")
    p(doc, "Ce mécanisme est la première ligne de défense contre les injections. On ne fait pas confiance au client pour envoyer des données propres — c'est une règle de base qu'on a appris à prendre au sérieux après avoir vu des exemples d'APIs publiques qui ne le font pas.")
    doc.add_page_break()

    # ── CHAPITRE 6 ─────────────────────────────────────────────
    h1(doc, "6. Stratégie de tests et validation qualité")
    h2(doc, "6.1. Tests unitaires : isoler pour vraiment tester")
    p(doc, "L'architecture asynchrone et la dépendance à des services externes (APIs, bases de données) compliquent la stratégie de test. Un test unitaire qui fait de vrais appels réseau n'est plus unitaire — il teste aussi la disponibilité du service externe, sa latence, et son état courant. Ce n'est pas reproductible, et ça rend les tests inutilisables en CI.")
    p(doc, "La solution est le mocking : on remplace les dépendances externes par des doublures qui retournent des fixtures prédéfinies. Pour les fonctions ETL, les réponses HTTP des APIs sont interceptées et substituées par des JSON statiques représentatifs de différents cas (réponse nominale, réponse partielle, erreur 503). Le framework utilisé est Pytest avec les plugins pytest-asyncio (pour les coroutines) et pytest-mock (pour le mocking).")

    h2(doc, "6.2. Tests d'intégration et tests de charge")
    p(doc, "Au-delà des tests unitaires, on a écrit des tests d'intégration qui testent des chemins complets à travers le système. Ces tests sont plus lents mais ils détectent des classes de bugs invisibles aux tests unitaires — notamment les problèmes de sérialisation/désérialisation entre services, et les effets de bord sur l'état des bases de données. Ils sont organisés dans le répertoire tests/integration/ et utilisent le TestClient de FastAPI pour simuler de vraies requêtes HTTP sans lancer le serveur.")
    p(doc, "Pour les tests de charge, on a utilisé Locust, un framework Python qui simule des utilisateurs concurrents. Le scénario de référence simule 500 utilisateurs requêtant la heatmap géospatiale toutes les secondes — ce qui représente une charge bien supérieure aux usages réalistes en phase de projet académique, mais permet de valider les choix d'architecture sous pression.")
    p(doc, "Les résultats : latence médiane (P50) à 142 ms, P95 à 380 ms, aucune erreur 5xx. La majorité des requêtes sont servies depuis le cache Redis, ce qui explique la faible latence médiane. Le P95 monte quand les requêtes tombent sur des données non cachées et doivent aller chercher dans PostgreSQL. Ce comportement est attendu et acceptable.")
    doc.add_page_break()

    # ── CHAPITRE 7 ─────────────────────────────────────────────
    h1(doc, "7. Interface client : React et rendu cartographique")
    h2(doc, "7.1. Architecture de la Single Page Application")
    p(doc, "L'interface est une application React construite avec Vite comme bundler. On a délibérément évité les frameworks « full-stack » type Next.js — le rendu côté serveur n'avait pas d'intérêt pour ce projet, et la complexité additionnelle non plus. Vite offre un serveur de développement très rapide (HMR quasi-instantané) et un build de production optimisé sans configuration complexe.")
    p(doc, "L'architecture des composants suit une séparation claire entre la présentation et la logique de données. Les composants de présentation sont stateless : ils reçoivent leurs données en props et ne gèrent pas d'état. Les conteneurs gèrent les appels API et l'état local via des hooks personnalisés. Cette séparation facilite les tests et rend le code plus lisible, mais elle demande une discipline qu'on n'a pas toujours réussi à maintenir sous pression — quelques composants ont fini par accumuler plus de responsabilités que prévu.")

    h2(doc, "7.2. Moteur cartographique et optimisation des performances")
    p(doc, "La visualisation cartographique s'appuie sur Leaflet.js, qui reste la référence pour les cartes interactives en web sans dépendance à un service tiers payant. On superpose des couches vectorielles GeoJSON sur des tuiles raster OpenStreetMap pour générer la heatmap de trafic.")
    p(doc, "Le défi technique principal était la performance de rendu. L'API retourne jusqu'à un millier de géométries GeoJSON représentant les segments de route avec leur état de trafic. React, par défaut, re-rend les composants enfants à chaque mise à jour de l'état parent. Dans notre cas, un rafraîchissement des données déclenchait un re-rendu complet de toutes les géométries — avec un drop de framerate mesurable pendant le re-rendu.")
    p(doc, "La solution a été une combinaison de mémoïsation : React.memo sur les composants de couche cartographique, useMemo pour les transformations de données GeoJSON, et useCallback pour les handlers d'événements. Avec ces optimisations, le re-rendu des données est devenu imperceptible à l'œil nu. C'est un cas d'usage typique où les outils de profiling du navigateur (React DevTools Profiler) sont indispensables — sans mesure, on ne saurait pas où optimiser.")
    doc.add_page_break()

    # ── CHAPITRE 8 ─────────────────────────────────────────────
    h1(doc, "8. Conformité RGPD et Privacy by Design")
    h2(doc, "8.1. Périmètre des données personnelles collectées")
    p(doc, "Le module de remontées citoyennes permet aux utilisateurs de signaler des incidents sur le réseau routier. Ces signalements contiennent deux catégories de données potentiellement personnelles : l'adresse IP de l'utilisateur (un identifiant réseau indirect) et la localisation GPS précise du signalement (qui peut permettre l'identification d'une personne si elle signale depuis son domicile ou son lieu de travail).")
    p(doc, "L'article 25 du RGPD impose que la protection des données personnelles soit intégrée dès la conception du système, pas ajoutée après coup. C'est le principe de Privacy by Design. En pratique, cela signifie qu'on ne doit jamais écrire de données personnelles brutes en base, même temporairement.")

    h2(doc, "8.2. Anonymisation de l'adresse IP par hachage salé")
    # CORRECTION: sel fixe via env variable, pas aléatoire par requête via secrets
    p(doc, "L'adresse IP source est hachée de façon irréversible avec SHA-256 avant toute persistance. SHA-256 seul ne suffit pas : les adresses IP sont peu nombreuses et une attaque par dictionnaire serait triviale. On ajoute donc un sel applicatif stocké dans la variable d'environnement GDPR_HASH_SALT, concaténé à l'adresse IP avant le hachage. Ce sel est commun à l'application et doit être changé en production (une valeur par défaut de développement est définie dans le code, clairement marquée comme à remplacer en prod).")
    p(doc, "Le résultat est un identifiant de corrélation qui permet de détecter des signalements répétés depuis la même IP sans jamais pouvoir retrouver l'adresse IP d'origine. C'est exactement le niveau de pseudonymisation adapté à ce type de données.")

    h2(doc, "8.3. Bruitage géographique des coordonnées GPS")
    p(doc, "La localisation GPS brute est trop précise pour être stockée telle quelle — elle permettrait de localiser un signalement à quelques mètres près, ce qui est suffisant pour identifier le domicile d'une personne dans un immeuble. On applique un décalage aléatoire uniforme sur les coordonnées de ±0.0015 degrés (environ ±150 mètres en Île-de-France), via la fonction apply_geo_blur() du router crowdsourcing.")
    p(doc, "Ce niveau de floutage est un compromis : il suffit à détruire l'information d'identification individuelle (on ne peut plus savoir si le signalement vient du bâtiment A ou du bâtiment B dans une rue), tout en préservant l'information macroscopique utile pour la modélisation du trafic (le signalement vient bien du quartier X, pas du quartier Y).")
    p(doc, "Les coordonnées floutées sont stockées directement dans le champ geometry via ST_MakePoint() de PostGIS. Les coordonnées brutes ne transitent jamais vers PostgreSQL. Un TTL (Time-To-Live) de 30 jours est appliqué à tous les signalements via le champ expires_at. L'API de lecture filtre automatiquement les enregistrements expirés (WHERE expires_at > NOW()), garantissant le droit à l'effacement prévu par l'article 17 du RGPD.")
    doc.add_page_break()

    # ── CHAPITRE 9 ─────────────────────────────────────────────
    h1(doc, "9. Empreinte environnementale et projections cloud")
    h2(doc, "9.1. Mesure de la consommation énergétique")
    p(doc, "L'entraînement de modèles de deep learning a une empreinte carbone non négligeable. On a intégré la librairie CodeCarbon dans nos scripts d'entraînement pour mesurer la consommation réelle. CodeCarbon interroge les compteurs d'énergie hardware (Intel RAPL pour le CPU et la RAM) et applique le facteur d'émission carbone de la grille électrique locale — environ 50 gCO2eq/kWh pour la France, qui est l'un des plus faibles d'Europe grâce au mix nucléaire. Les émissions sont loggées en kg CO2 équivalent et sauvegardées dans le répertoire logs/carbon/.")
    p(doc, "Cette mesure a motivé deux décisions pratiques : l'early stopping (qui coupe les époques d'entraînement dès que la val_loss stagne, évitant des centaines de passes inutiles), et la planification des réentraînements en dehors des heures de pointe électriques. C'est une optimisation marginale à notre échelle, mais qui traduit une démarche Green IT cohérente et mesurable.")

    h2(doc, "9.2. Modélisation du coût cloud (FinOps)")
    p(doc, "Si le projet devait passer en production sur un cloud public (GCP ou AWS), l'estimation du coût mensuel récurrent s'établit autour de 480 euros HT. Ce chiffre est construit de façon bottom-up à partir des composants effectivement utilisés.")
    p(doc, "Les deux instances de calcul backend en load balancing représentent environ 80 euros. Le cluster PostgreSQL/PostGIS managé (RDS sur AWS ou Cloud SQL sur GCP) est le poste le plus cher à l'usage — environ 150 euros pour une instance avec réplication et backups automatiques. Le nœud ElastiCache Redis s'établit autour de 40 euros. L'instance GPU pour les réentraînements du LSTM (une Nvidia T4 en mode spot ou préemptible, allouée uniquement pendant les fenêtres de réentraînement nocturnes) représente 200 euros environ. Le stockage objet (type S3 ou MinIO) pour les artefacts de modèles représente un poste marginal d'environ 10 euros.")
    p(doc, "En inférence, le LSTM tourne correctement sur CPU avec une latence acceptable. En réentraînement, le GPU divise le temps par un facteur 8 à 10. Pour un système où les réentraînements sont hebdomadaires et nocturnes, l'utilisation d'instances GPU spot serait suffisante.")
    doc.add_page_break()

    # ── CHAPITRE 10 ────────────────────────────────────────────
    h1(doc, "10. Limites du système et perspectives")
    h2(doc, "10.1. La limite fondamentale : l'absence de topologie")
    p(doc, "La lacune technique la plus sérieuse du système actuel est que chaque capteur est traité comme un capteur isolé. Le modèle ne sait pas que le capteur A est en amont du capteur B sur le même axe, et que la congestion au capteur A va mécaniquement se propager vers B quelques minutes plus tard.")
    p(doc, "En termes de théorie des graphes, le réseau routier est un graphe orienté valué, et la propagation de la congestion obéit à des lois qui ressemblent à la mécanique des fluides — avec des effets de compression, de contournement, et d'onde de choc. Ignorer cette structure, c'est ignorer une source d'information prédictive considérable.")
    p(doc, "L'axe d'amélioration prioritaire est le remplacement (ou le complément) des modèles actuels par des réseaux de neurones sur graphes spatio-temporels (Spatial-Temporal Graph Convolutional Networks, STGCN). Ces architectures prennent en entrée non seulement les séries temporelles de chaque capteur, mais aussi la matrice d'adjacence du réseau routier. Les papiers de référence (DCRNN de Li et al., 2018, et Graph WaveNet de Wu et al., 2019) montrent des améliorations significatives par rapport aux approches univariées sur des benchmarks standards.")
    p(doc, "Implémenter un STGCN aurait cependant demandé un temps de développement et de mise au point bien supérieur à ce que notre calendrier permettait. On a fait le choix pragmatique de construire quelque chose de fonctionnel avec les approches hybrides SARIMAX/LSTM, quitte à documenter honnêtement les limites.")

    h2(doc, "10.2. Industrialisation et MLOps")
    p(doc, "Le pipeline de modélisation actuel est fonctionnel mais non industrialisé. Un réentraînement implique de lancer manuellement un script, de vérifier les métriques à la main, et de copier les artefacts manuellement. C'est acceptable pour un prototype académique, ça ne l'est pas pour un système en production.")
    p(doc, "Une stack MLOps complète résoudrait ces problèmes. MLFlow permettrait de tracker les expériences, versionner les modèles, et comparer les métriques entre les différentes runs d'entraînement. Un pipeline de Continuous Training détecterait automatiquement la dérive statistique des données d'entrée (data drift) et déclencherait un réentraînement en conséquence.")
    p(doc, "La surveillance du data drift est particulièrement importante pour un système de trafic. Les habitudes de mobilité changent — la généralisation du télétravail après 2020 a significativement modifié les patterns de déplacement. Un modèle entraîné sur des données pré-COVID prédirait mal le trafic actuel. Le monitoring en production permet de détecter ce phénomène avant qu'il ne dégrade silencieusement les performances.")

    h2(doc, "10.3. Sources de données complémentaires")
    p(doc, "Le système actuel utilise trois sources principales. Plusieurs sources supplémentaires pourraient améliorer la précision des prédictions.")
    p(doc, "Les données d'événements (concerts, manifestations, marchés, matchs) créent des perturbations localisées prévisibles. Ces événements sont connus à l'avance et pourraient être intégrés comme variables exogènes dans le SARIMAX.")
    p(doc, "Les données de travaux et de fermetures de voirie sont disponibles via la base BORA (Base des Obstacles de la Voirie d'Île-de-France). L'impact des travaux sur le trafic est très localisant et durable — un chantier peut modifier les patterns pendant plusieurs mois.")
    p(doc, "Enfin, les données de flottes de véhicules connectés (GPS flottants) offrent une couverture spatiale bien supérieure aux boucles de comptage fixes. L'accès à ces données implique des accords avec les fournisseurs et des considérations RGPD supplémentaires, mais c'est la direction vers laquelle vont les systèmes de trafic de nouvelle génération.")
    doc.add_page_break()

    # ── CONCLUSION ─────────────────────────────────────────────
    h1(doc, "Conclusion")
    p(doc, "UrbanFlow est une plateforme fonctionnelle de prédiction de trafic en temps quasi-réel. Elle démontre la faisabilité d'une architecture distribuée complète — de l'ingestion des données brutes jusqu'à la restitution cartographique — avec des contraintes réalistes de latence, de sécurité et de conformité réglementaire.")
    p(doc, "Sur le plan technique, les points qu'on retient comme contributions réelles : l'implémentation de l'algorithme d'arbitrage dynamique entre SARIMAX et LSTM, qui améliore la précision de façon consistante sur les deux régimes de trafic ; l'architecture de persistance polyglotte Redis/PostgreSQL qui satisfait à la fois les contraintes de latence temps réel et les besoins de stockage historique ; et la chaîne complète d'anonymisation RGPD intégrée dès la conception.")
    p(doc, "Les limites sont réelles et documentées honnêtement : l'absence de modélisation topologique du réseau est la lacune principale, et elle est structurelle par rapport aux modèles utilisés. La non-industrialisation du cycle de vie ML est un second point qui rendrait le système fragile en production réelle.")
    p(doc, "Ce que ce projet nous a appris, c'est que la difficulté en Data Engineering n'est pas dans les algorithmes eux-mêmes — ils sont documentés, les librairies existent, les tutoriels abondent. La difficulté est dans l'intégration : faire fonctionner ensemble des composants qui ont des contrats différents, des modes de défaillance différents, et des exigences de performance différentes. C'est exactement ce que ce Mastère nous avait dit préparer à faire, et on mesure maintenant concrètement ce que ça signifie.")
    p(doc, "Les perspectives identifiées — STGCN pour la modélisation topologique, stack MLOps pour l'industrialisation, intégration de sources d'événements — sont réalistes et forment une feuille de route cohérente. Elles supposent cependant des moyens supplémentaires, notamment pour l'accès à des données de trafic plus granulaires et à de l'infrastructure GPU en production.")
    doc.add_page_break()

    # ── RÉFÉRENCES ─────────────────────────────────────────────
    h1(doc, "Références et sources")

    h2(doc, "Données et APIs")
    p(doc, "–  API Open Data Voirie — data.gouv.fr / Direction des Routes d'Île-de-France")
    p(doc, "–  API PRIM (Prochains Passages Île-de-France Mobilités) — GTFS-RT")
    p(doc, "–  OpenWeatherMap API — Données météorologiques historiques et temps réel")

    h2(doc, "Librairies et frameworks")
    p(doc, "–  FastAPI — Sebastián Ramírez — https://fastapi.tiangolo.com")
    p(doc, "–  httpx — Encodé-Magnusson et al. — Client HTTP asynchrone Python")
    p(doc, "–  asyncpg — MagicStack — PostgreSQL async driver")
    p(doc, "–  PyTorch / Keras (TensorFlow) — Framework deep learning (implémentation LSTM)")
    p(doc, "–  statsmodels — Implémentation SARIMAX")
    p(doc, "–  Tenacity — Bibliothèque de résilience (retry/exponential backoff)")
    p(doc, "–  CodeCarbon — Mesure de l'empreinte carbone des calculs")
    p(doc, "–  React + Vite — Frontend SPA")
    p(doc, "–  Leaflet.js — Rendu cartographique")
    p(doc, "–  Locust — Tests de charge")

    h2(doc, "Références académiques")
    p(doc, "–  Li, Y. et al. (2018). Diffusion Convolutional Recurrent Neural Network: Data-Driven Traffic Forecasting. ICLR 2018.")
    p(doc, "–  Wu, Z. et al. (2019). Graph WaveNet for Deep Spatial-Temporal Graph Modeling. IJCAI 2019.")
    p(doc, "–  Box, G.E.P., Jenkins, G.M. (1976). Time Series Analysis: Forecasting and Control. Holden-Day.")
    p(doc, "–  Hochreiter, S., Schmidhuber, J. (1997). Long Short-Term Memory. Neural Computation, 9(8).")
    p(doc, "–  Hyndman, R.J., Athanasopoulos, G. (2021). Forecasting: Principles and Practice. 3rd ed. OTexts.")

    doc.save("Rapport_Technique_UrbanFlow_CORRIGE.docx")
    print("Rapport généré : Rapport_Technique_UrbanFlow_CORRIGE.docx")

if __name__ == "__main__":
    create()
